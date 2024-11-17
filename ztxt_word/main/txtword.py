from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from io import BytesIO

# Константы для размеров изображения (в дюймах)
MARGIN_INCH = 0.17322835
IMAGE_WIDTH = 5.314961
IMAGE_HEIGHT = 8.16929134

def process_text_to_word(
    text, 
    font_path=r"main/static/fonts/Salavat-Font4You.ttf", 
    font_size=30, 
    notebook_image_path=r"main/static/img/plate.jpg",
    padding_left_odd=65, 
    padding_right_odd=45, 
    padding_left_even=45, 
    padding_right_even=45, 
    padding_top=35, 
    padding_bottom=10, 
    line_spacing=20,
    justify_threshold=0.8  # Порог для выравнивания строки (80% ширины страницы)
):
    """Обрабатывает текст и возвращает Word-файл как бинарный объект BytesIO с условием выравнивания текста по ширине."""

    def split_text_into_lines(text, font, max_width):
        lines = text.splitlines()
        final_lines = []
        for line in lines:
            words = line.split(' ')
            current_line = ""
            for word in words:
                test_line = current_line + word + " "
                if font.getbbox(test_line)[2] <= max_width:
                    current_line = test_line
                else:
                    final_lines.append(current_line.rstrip())
                    current_line = word + " "
            final_lines.append(current_line.rstrip())
        return final_lines

    def draw_justified_text(draw, text, font, x, y, max_width, justify_threshold):
        words = text.split()
        line_width = sum(font.getbbox(word)[2] for word in words) + font.getbbox(" ")[2] * (len(words) - 1)
        
        # Если ширина строки меньше порогового значения, рисуем текст с обычными пробелами
        if line_width < max_width * justify_threshold or len(words) == 1:
            draw.text((x, y), text, font=font, fill=(0, 0, 0, 255))
            return
        
        # Иначе распределяем пробелы равномерно
        total_spacing = max_width - line_width + font.getbbox(" ")[2] * (len(words) - 1)
        space_width = total_spacing // (len(words) - 1)

        cursor_x = x
        for i, word in enumerate(words):
            draw.text((cursor_x, y), word, font=font, fill=(0, 0, 0, 255))
            cursor_x += font.getbbox(word)[2] + (space_width if i < len(words) - 1 else 0)

    font = ImageFont.truetype(font_path, font_size)
    notebook_image_odd = Image.open(notebook_image_path)
    notebook_image_even = notebook_image_odd.transpose(Image.FLIP_LEFT_RIGHT)
    
    max_height = notebook_image_odd.height - (padding_top + padding_bottom)
    max_width_odd = notebook_image_odd.width - (padding_left_odd + padding_right_odd)
    lines = split_text_into_lines(text, font, max_width_odd)
    
    line_height = font.getbbox("A")[3] - font.getbbox("A")[1] + line_spacing
    max_lines_per_page = max_height // line_height
    pages = [lines[i:i + max_lines_per_page] for i in range(0, len(lines), max_lines_per_page)]
    images = []

    for page_number, page_lines in enumerate(pages):
        notebook_image = notebook_image_odd.copy() if page_number % 2 == 0 else notebook_image_even.copy()
        padding_left = padding_left_odd if page_number % 2 == 0 else padding_left_even
        padding_right = padding_right_odd if page_number % 2 == 0 else padding_right_even
        
        text_image = Image.new('RGBA', notebook_image.size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(text_image)
        y_position = padding_top
        for line in page_lines:
            draw_justified_text(draw, line, font, padding_left, y_position, max_width_odd, justify_threshold)
            y_position += line_height
        
        combined_image = Image.alpha_composite(notebook_image.convert('RGBA'), text_image)
        img_byte_arr = BytesIO()
        combined_image.save(img_byte_arr, format="PNG")
        images.append(img_byte_arr)

    # Создание Word документа в памяти
    doc_stream = BytesIO()
    document = Document()
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(MARGIN_INCH)

    for idx in range(0, len(images), 2):
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        
        # Добавляем изображения из BytesIO в документ
        run.add_picture(images[idx], width=Inches(IMAGE_WIDTH), height=Inches(IMAGE_HEIGHT))
        if idx + 1 < len(images):
            run.add_picture(images[idx + 1], width=Inches(IMAGE_WIDTH), height=Inches(IMAGE_HEIGHT))

    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream